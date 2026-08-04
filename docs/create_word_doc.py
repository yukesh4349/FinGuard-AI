import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_finguard_docx():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("FinGuard AI - System Master Documentation")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 22, 16) # #1A1610

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Main System Admin Login Credentials, PostgreSQL Schemas & App Manual\nRepository: https://github.com/yukesh4349/FinGuard-AI")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(138, 117, 88)

    doc.add_paragraph() # Spacing

    # Section 1: Main Admin Credentials
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. System Main Admin Credentials")
    r_h1.font.color.rgb = RGBColor(26, 22, 16)

    p1 = doc.add_paragraph()
    p1.add_run("As the System Owner, you have full governance and super-admin access to inspect all registered users, emails, mobile numbers, assigned roles, and passwords stored in the PostgreSQL database.").font.size = Pt(10.5)

    # Table of Admin Credentials
    t_admin = doc.add_table(rows=1, cols=4)
    t_admin.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t_admin.rows[0].cells
    headers = ["Field Name", "Credential Value", "Role Level", "Access Scope"]
    for i, head in enumerate(headers):
        hdr_cells[i].text = head
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Background color
        shading = parse_xml(r'<w:shd {} w:fill="1A1610"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    admin_data = [
        ("Main Admin Login ID", "admin@finguard.ai", "Super Admin", "Full User & DB Password Control"),
        ("Alternative Login ID", "ADMIN-001 or admin", "Super Admin", "Direct System Master Access"),
        ("Main Admin Password", "admin123 (or admin)", "Super Admin", "Unrestricted Master Password"),
        ("System Mobile Number", "9999999999", "Super Admin", "Root Alert Communications"),
    ]

    for row_data in admin_data:
        row_cells = t_admin.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph() # Spacing

    # Section 2: All Registered Users & Passwords Table
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. PostgreSQL Registered Users & Passwords Directory")
    r_h2.font.color.rgb = RGBColor(26, 22, 16)

    p2 = doc.add_paragraph()
    p2.add_run("Below is the complete database registry of pre-seeded and dynamically registered accounts in FinGuard AI:").font.size = Pt(10.5)

    t_users = doc.add_table(rows=1, cols=5)
    t_users.alignment = WD_TABLE_ALIGNMENT.CENTER
    u_hdr = t_users.rows[0].cells
    u_headers = ["User ID / Identifier", "Company / Store Name", "Email Address", "Password", "Mobile Number"]
    for i, head in enumerate(u_headers):
        u_hdr[i].text = head
        u_hdr[i].paragraphs[0].runs[0].font.bold = True
        u_hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="5C705E"/>'.format(nsdecls('w')))
        u_hdr[i]._tc.get_or_add_tcPr().append(shading)

    users_data = [
        ("ADMIN-001 (Super Admin)", "FinGuard System Governance", "admin@finguard.ai", "admin123", "9999999999"),
        ("OWNER-METRO-8492 (Business Owner)", "Metro Superstore Ltd", "owner@metrosuperstore.com", "FG-8924-XK9", "9876543210"),
        ("accountant@metrosuperstore.com (Store Accountant)", "Metro Superstore Ltd", "accountant@metrosuperstore.com", "FG-CA-2026", "9876523451"),
        ("cashier.billing@metrosuperstore.com (Cashier & Billing)", "Metro Superstore Ltd", "cashier.billing@metrosuperstore.com", "FG-BILL-789", "9876545673"),
        ("manager.stock@metrosuperstore.com (Stock Manager)", "Metro Superstore Ltd", "manager.stock@metrosuperstore.com", "FG-STOCK-552", "9876534562"),
    ]

    for row_data in users_data:
        row_cells = t_users.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph() # Spacing

    # Section 3: PostgreSQL Configuration & Schema
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. PostgreSQL Connection & Table Schema Setup")
    r_h3.font.color.rgb = RGBColor(26, 22, 16)

    p3 = doc.add_paragraph()
    p3.add_run("Database Name: ").bold = True
    p3.add_run("finguard_db\n")
    p3.add_run("Default Host: ").bold = True
    p3.add_run("localhost:5432\n")
    p3.add_run("Default User: ").bold = True
    p3.add_run("postgres\n")

    p3_code = doc.add_paragraph()
    p3_code.paragraph_format.left_indent = Inches(0.4)
    run_code = p3_code.add_run("""CREATE TABLE users (
    id SERIAL PRIMARY KEY,
    user_id VARCHAR(100) UNIQUE NOT NULL,
    company_name VARCHAR(255) NOT NULL,
    mobile_number VARCHAR(20) NOT NULL,
    email VARCHAR(255) UNIQUE NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    role VARCHAR(50) DEFAULT 'owner',
    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
);""")
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(9)

    doc.add_paragraph() # Spacing

    # Section 4: Key Application Features
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. FinGuard AI Feature Summary")
    r_h4.font.color.rgb = RGBColor(26, 22, 16)

    features = [
        "Simple English Everywhere: Completely removed complex jargon to make finance, taxes, and stock management simple for everyone.",
        "Login & Sign Up Image Sidebar: Left sidebar text replaced with high-tech FinGuard AI visual image banner.",
        "Privacy & Passwords: Password fields allow user-defined passwords without showing or auto-filling credentials.",
        "Numeric-Only Mobile Fields: Restricted mobile phone inputs strictly to digits 0-9.",
        "Restructured Owner Dashboard: Displays Profit/Loss and Revenue graphs first, floating AI Chatbot in bottom-right corner, Fraud Alerts as live chat history on the right, action buttons & stock/supplier/employee detail cards on the left, and cashflow history at the bottom.",
    ]

    for feat in features:
        p_f = doc.add_paragraph(style='List Bullet')
        p_f.add_run(feat).font.size = Pt(10)

    # Save document
    doc.save("docs/FinGuard_AI_System_Credentials_and_Documentation.docx")
    print("Word document created successfully: docs/FinGuard_AI_System_Credentials_and_Documentation.docx")

if __name__ == "__main__":
    create_finguard_docx()
